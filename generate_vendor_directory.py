#!/usr/bin/env python3
"""
Generate 2026 Disability Community Resource Fair Vendor Directory

This script:
1. Reads 2026 Vendor Requests from fair_logistics.xlsx
2. Extracts vendor requests CSV
3. Reads vendor details from _posts markdown files
4. Combines vendor request numbers with detailed vendor information
5. Generates a comprehensive vendor directory markdown with:
   - Vendor listings organized by table number
   - Full service descriptions (no truncation)
   - Service category index (10 main categories)
   - Age group index (5 age ranges)
"""

import os
import re
import argparse
import frontmatter
import csv
import zipfile
import xml.etree.ElementTree as ET
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

EVENT_DATE_TIME = 'Saturday, May 2, 2026 from 10:00 AM to 1:00 PM'

def read_vendor_filename_map(map_path):
    """Read a manual vendor filename / override mapping from CSV in _data.

    Expected CSV format (header optional):
    request_name,filename,table_number_override
    """
    mapping = {}
    if not os.path.exists(map_path):
        return mapping

    try:
        with open(map_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                if not row:
                    continue
                # allow header row detection
                if row[0].strip().lower() in ('request_name', 'vendor'):
                    continue

                req = row[0].strip()
                if not req:
                    continue
                filename = row[1].strip() if len(row) > 1 else ''
                override = row[2].strip() if len(row) > 2 else ''
                mapping[normalize_name(req)] = {
                    'request_name': req,
                    'filename': filename,
                    'table_override': override,
                }
    except Exception:
        return mapping

    return mapping

def extract_vendors_from_excel(excel_path, output_csv_path):
    """Extract vendors from fair_logistics.xlsx 2026 Vendor Requests sheet."""
    
    def extract_cell_value(cell, shared_strings):
        """Extract value from a cell, handling shared strings."""
        value_elem = cell.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v')
        
        if value_elem is not None:
            cell_type = cell.get('t')
            text_value = value_elem.text or ''
            
            if cell_type == 's':
                try:
                    idx = int(text_value)
                    return shared_strings[idx] if idx < len(shared_strings) else ''
                except:
                    return text_value
            else:
                return text_value
        
        rich_text = cell.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}is')
        if rich_text is not None:
            text_parts = rich_text.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t')
            return ''.join([t.text or '' for t in text_parts])
        
        return ''
    
    with zipfile.ZipFile(excel_path, 'r') as zip_ref:
        # Load shared strings
        shared_strings = []
        try:
            shared_strings_xml = zip_ref.read('xl/sharedStrings.xml')
            ss_root = ET.fromstring(shared_strings_xml)
            string_items = ss_root.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}si')
            
            for si in string_items:
                t_elem = si.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t')
                if t_elem is not None:
                    shared_strings.append(t_elem.text or '')
                else:
                    text_parts = si.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}r/{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t')
                    shared_strings.append(''.join([t.text or '' for t in text_parts]))
        except Exception as e:
            print(f"Warning: Error loading shared strings: {e}")
        
        # Read the 2026 Vendor Requests sheet
        worksheet_xml = zip_ref.read('xl/worksheets/sheet10.xml')
        ws_root = ET.fromstring(worksheet_xml)
        
        rows = ws_root.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row')
        
        all_data = []
        for row in rows:
            cells = row.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c')
            row_data = []
            for cell in cells:
                value = extract_cell_value(cell, shared_strings)
                row_data.append(value)
            if any(row_data):
                all_data.append(row_data)
        
        # Save to CSV
        with open(output_csv_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerows(all_data)
        
        return len(all_data)

def read_vendor_requests(csv_path):
    """Read vendor requests from CSV."""
    vendor_requests = {}
    
    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        for row in reader:
            if len(row) >= 2 and row[0].strip() and row[1].strip():
                vendor_num = row[0].strip()
                vendor_name = row[1].strip()
                
                # Skip header rows that may contain 'VENDOR REQUESTS' or similar
                if 'VENDOR REQUESTS' in vendor_num.upper() or vendor_name.strip().upper() == 'ARRIVED':
                    continue
                # Preserve the original table identifier (could be a range like '19-21', 'Lobby', or '3 & 4')
                # Store as string so we don't drop valid entries that are non-integer.
                vendor_requests[vendor_name] = vendor_num
    
    return vendor_requests

def read_vendor_details(posts_dir):
    """Read vendor details from markdown files in _posts."""
    vendor_details = {}
    
    md_files = sorted([f for f in os.listdir(posts_dir) if f.endswith('.md')])
    
    for md_file in md_files:
        filepath = os.path.join(posts_dir, md_file)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                post = frontmatter.load(f)
            
            vendor_name = md_file.split('-', 3)[-1].replace('.md', '')
            
            metadata = {
                'title': post.metadata.get('title', vendor_name),
                'categories': post.metadata.get('categories', []),
                'tags': post.metadata.get('tags', []),
                'content': post.content,
                'filename': md_file
            }
            
            vendor_details[vendor_name] = metadata
        except Exception as e:
            pass
    
    return vendor_details

def read_vendor_map_entries(vendor_map_path):
    """Read vendor labels from the vendor map workbook drawing layer."""
    vendor_map_entries = {}

    if not os.path.exists(vendor_map_path):
        return []

    drawing_path = 'xl/drawings/drawing1.xml'
    ns = {
        'xdr': 'http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    }

    with zipfile.ZipFile(vendor_map_path, 'r') as zip_ref:
        if drawing_path not in zip_ref.namelist():
            return []

        root = ET.fromstring(zip_ref.read(drawing_path))
        for shape in root.findall('.//xdr:sp', ns):
            text_parts = [t.text or '' for t in shape.findall('.//a:t', ns)]
            raw_text = ''.join(text_parts).strip()
            if not raw_text:
                continue

            match = re.match(r'^(Lobby|\d+\s*&\s*\d+|\d+(?:-\d+)?)\s+(.*\S)\s*$', raw_text)
            if not match:
                continue

            number_text = match.group(1).replace(' ', '')
            display_name = match.group(2).strip()
            normalized_name = normalize_name(display_name)

            if not display_name or not normalized_name:
                continue

            entry = vendor_map_entries.get(normalized_name)
            if entry is None:
                entry = {
                    'display_name': display_name,
                    'normalized_name': normalized_name,
                    'number_values': set(),
                }
                vendor_map_entries[normalized_name] = entry

            entry['number_values'].update(parse_table_number_set(number_text))

    return list(vendor_map_entries.values())

def clean_category_name(category):
    """Clean category name by replacing underscores with spaces."""
    return category.replace('_', ' ').strip()

def parse_categories(categories_str):
    """Parse categories string into individual category names.
    
    Categories are SPACE-delimited (not comma-delimited).
    Within each category, replace underscores with spaces.
    Commas within categories are kept as-is.
    """
    if isinstance(categories_str, str):
        # Split by spaces to get individual categories
        cats = [c.strip() for c in categories_str.split()]
    else:
        cats = categories_str if isinstance(categories_str, list) else []
    
    # Clean up each category
    cleaned = []
    for cat in cats:
        if cat:
            # Replace underscores with spaces within the category
            cat = cat.replace('_', ' ')
            if cat:
                cleaned.append(cat)
    
    return cleaned

def parse_tags(tags_str):
    """Parse tags string into individual tag names."""
    if isinstance(tags_str, str):
        tags = [t.strip() for t in tags_str.split()]
    else:
        tags = tags_str if isinstance(tags_str, list) else []
    
    cleaned = []
    for tag in tags:
        if tag:
            tag = tag.replace('_', ' ')
            if tag:
                cleaned.append(tag)
    
    return cleaned

def normalize_name(s):
    if not s:
        return ''
    ns = s.lower().replace('&', 'and')
    ns = re.sub(r'[^a-z0-9\s]', '', ns)
    return ' '.join(ns.split())

def names_match(left, right, threshold=0.85):
    left_norm = normalize_name(left)
    right_norm = normalize_name(right)
    if not left_norm or not right_norm:
        return False
    if left_norm == right_norm:
        return True
    if left_norm in right_norm or right_norm in left_norm:
        return True

    import difflib

    return difflib.SequenceMatcher(a=left_norm, b=right_norm).ratio() >= threshold

def parse_table_number_set(number_text):
    if number_text is None:
        return set()

    text = str(number_text).strip()
    if not text:
        return set()

    if text.lower() == 'lobby':
        return {'Lobby'}

    normalized = text.replace(' ', '')
    parts = re.split(r'(?:&|,|/|;|\band\b)', normalized, flags=re.IGNORECASE)
    numbers = set()

    for part in parts:
        token = part.strip()
        if not token:
            continue

        range_match = re.fullmatch(r'(\d+)-(\d+)', token)
        if range_match:
            start = int(range_match.group(1))
            end = int(range_match.group(2))
            if start <= end:
                numbers.update(range(start, end + 1))
            else:
                numbers.update(range(end, start + 1))
            continue

        if token.isdigit():
            numbers.add(int(token))
            continue

        numbers.add(token)

    return numbers

def format_table_number_set(number_values):
    if not number_values:
        return ''

    def sort_key(value):
        if isinstance(value, int):
            return (0, value)
        return (1, str(value))

    return ', '.join(str(value) for value in sorted(number_values, key=sort_key))

def extract_docx_template_styles(template_path):
    """Extract font, size, and color information from an existing DOCX template.
    
    Returns a dictionary with style information for replication.
    """
    if not os.path.exists(template_path):
        return {}
    
    try:
        template_doc = Document(template_path)
        extracted = {}
        
        # Extract from existing character styles (tag styles)
        for style_name in ['Vendor Service Tag', 'Vendor Age Tag', 'Vendor All Ages Tag']:
            if style_name in template_doc.styles:
                style = template_doc.styles[style_name]
                font = style.font
                extracted[style_name] = {
                    'font_name': font.name,
                    'font_size': font.size,
                    'color': font.color.rgb if font.color and font.color.rgb else None,
                    'bold': font.bold,
                }
        
        # Extract from heading styles
        for level in [1, 2, 3]:
            style_name = f'Heading {level}'
            if style_name in template_doc.styles:
                style = template_doc.styles[style_name]
                font = style.font
                extracted[style_name] = {
                    'font_name': font.name,
                    'font_size': font.size,
                    'color': font.color.rgb if font.color and font.color.rgb else None,
                    'bold': font.bold,
                }
        
        # Extract from normal paragraph style
        if 'Normal' in template_doc.styles:
            style = template_doc.styles['Normal']
            font = style.font
            extracted['Normal'] = {
                'font_name': font.name,
                'font_size': font.size,
                'color': font.color.rgb if font.color and font.color.rgb else None,
            }
        
        return extracted
    except Exception as e:
        print(f"Warning: Could not extract styles from template: {e}")
        return {}


def ensure_docx_tag_styles(document, extracted_styles=None):
    """Create reusable character styles for DOCX tag rendering.
    
    If extracted_styles is provided, uses those colors/fonts instead of defaults.
    """
    # Default colors (fallback)
    default_specs = [
        ('Vendor Service Tag', RGBColor(2, 61, 102)),
        ('Vendor Age Tag', RGBColor(74, 110, 27)),
        ('Vendor All Ages Tag', RGBColor(74, 23, 102)),
    ]

    for style_name, default_color in default_specs:
        if style_name not in document.styles:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
            
            # Use extracted style if available, otherwise use defaults
            if extracted_styles and style_name in extracted_styles:
                extracted = extracted_styles[style_name]
                style.font.size = extracted.get('font_size') or Pt(9)
                style.font.bold = extracted.get('bold', True)
                if extracted.get('color'):
                    style.font.color.rgb = extracted.get('color')
                else:
                    style.font.color.rgb = default_color
                if extracted.get('font_name'):
                    style.font.name = extracted.get('font_name')
            else:
                style.font.size = Pt(9)
                style.font.bold = True
                style.font.color.rgb = default_color


def add_docx_tag_run(paragraph, text, style_name):
    run = paragraph.add_run(text)
    run.style = style_name
    return run


def add_docx_badge_paragraph(document, service_labels=None, age_labels=None, all_ages=False):
    """Add a single-line badge paragraph to a docx document."""
    paragraph = document.add_paragraph()

    labels_added = False
    for label in service_labels or []:
        if label:
            if labels_added:
                paragraph.add_run(' ')
            add_docx_tag_run(paragraph, label, 'Vendor Service Tag')
            labels_added = True

    if all_ages:
        if labels_added:
            paragraph.add_run(' ')
        add_docx_tag_run(paragraph, 'All ages', 'Vendor All Ages Tag')
        return paragraph

    for label in age_labels or []:
        if label:
            if labels_added:
                paragraph.add_run(' ')
            add_docx_tag_run(paragraph, label, 'Vendor Age Tag')
            labels_added = True

    return paragraph

def add_docx_paragraphs(document, content):
    for block in content.split('\n\n'):
        block = block.strip()
        if block:
            document.add_paragraph(block)

def build_vendor_directory_banner_markdown(logo_href):
    return (
        f'<p align="center"><img src="{logo_href}" alt="Disability Community Resource Fair logo" style="max-width:220px;height:auto;"></p>\n\n'
        f'**Event Date & Time:** {EVENT_DATE_TIME}\n\n'
    )

def build_vendor_directory_banner_html(logo_href):
    return [
        f'<div style="text-align:center;margin:1rem 0 1.25rem 0;">',
        f'<img src="{logo_href}" alt="Disability Community Resource Fair logo" style="max-width:220px;height:auto;display:block;margin:0 auto 0.75rem auto;">',
        f'<div style="font-weight:600;">{EVENT_DATE_TIME}</div>',
        '</div>',
    ]

def add_vendor_directory_banner_docx(document, logo_path):
    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(2.4))

    event_paragraph = document.add_paragraph(EVENT_DATE_TIME)
    event_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    event_paragraph.runs[0].bold = True

def write_vendor_directory_docx(output_path, active, cancelled_entries, all_tags, num_to_name, age_groups):
    """Write the vendor directory as a DOCX document.
    
    Extracts custom styles from existing DOCX (if it exists) and applies them to the new output.
    """
    # Extract styles from existing DOCX template if it exists
    extracted_styles = extract_docx_template_styles(output_path)
    
    document = Document()
    ensure_docx_tag_styles(document, extracted_styles=extracted_styles)
    
    # Apply extracted heading styles if available
    if extracted_styles:
        for level in [1, 2, 3]:
            heading_style = f'Heading {level}'
            if heading_style in extracted_styles and heading_style in document.styles:
                style_info = extracted_styles[heading_style]
                style = document.styles[heading_style]
                if style_info.get('font_size'):
                    style.font.size = style_info.get('font_size')
                if style_info.get('font_name'):
                    style.font.name = style_info.get('font_name')
                if style_info.get('color'):
                    style.font.color.rgb = style_info.get('color')
                # Explicitly set headings to not bold
                style.font.bold = False
        
        # Apply extracted normal style
        if 'Normal' in extracted_styles and 'Normal' in document.styles:
            style_info = extracted_styles['Normal']
            style = document.styles['Normal']
            if style_info.get('font_size'):
                style.font.size = style_info.get('font_size')
            if style_info.get('font_name'):
                style.font.name = style_info.get('font_name')
            if style_info.get('color'):
                style.font.color.rgb = style_info.get('color')
    
    document.core_properties.title = 'Disability Community Resource Fair 2026 - Vendor Directory'

    logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(output_path))), 'assets', 'img', 'logo.png')
    add_vendor_directory_banner_docx(document, logo_path)

    document.add_heading('Disability Community Resource Fair 2026 - Vendor Directory', level=1)
    document.add_paragraph(
        'This is the comprehensive vendor directory for the 2026 Disability Community Resource Fair. '
        'Vendors are listed by table number, with an index organized by age groups served.'
    )

    document.add_heading('Vendor Listings by Table Number', level=2)

    for vendor in active:
        title = vendor['detail']['title'] if vendor['detail'] else vendor['request_name']
        document.add_heading(f"{vendor['number']}. {title}", level=3)

        if vendor['detail']:
            categories = parse_categories(vendor['detail']['categories'])
            tags = parse_tags(vendor['detail']['tags'])
            tag_set = set(tags)

            if categories or tags:
                add_docx_badge_paragraph(
                    document,
                    service_labels=categories,
                    age_labels=None if all(a in tag_set for a in age_groups) else tags,
                    all_ages=bool(tags) and all(a in tag_set for a in age_groups),
                )

            content = vendor['detail']['content'].strip()
            if content:
                add_docx_paragraphs(document, content)

    if cancelled_entries:
        document.add_heading('Unable to Attend', level=2)
        for vendor in cancelled_entries:
            title = vendor['detail']['title'] if vendor['detail'] else vendor['request_name']
            document.add_heading(title, level=3)
            if vendor['detail']:
                categories = parse_categories(vendor['detail']['categories'])
                tags = parse_tags(vendor['detail']['tags'])
                tag_set = set(tags)

                if categories or tags:
                    add_docx_badge_paragraph(
                        document,
                        service_labels=categories,
                        age_labels=None if all(a in tag_set for a in age_groups) else tags,
                        all_ages=bool(tags) and all(a in tag_set for a in age_groups),
                    )

                content = vendor['detail']['content'].strip()
                if content:
                    add_docx_paragraphs(document, content)
            else:
                document.add_paragraph('Description unavailable.')

    document.add_heading('Index by Age Group Served', level=2)
    
    # Build a combined index with both active and cancelled vendors
    # For cancelled vendors, we'll use their request_name since they don't have numbers
    combined_all_tags = defaultdict(set)
    cancelled_names = set()
    
    # Add active vendors to combined index
    for tag, vendors in all_tags.items():
        for vendor_num in vendors:
            combined_all_tags[tag].add(('number', vendor_num))
    
    # Add cancelled vendors to combined index
    for cancelled_vendor in cancelled_entries:
        if cancelled_vendor['detail']:
            tags = parse_tags(cancelled_vendor['detail'].get('tags', ''))
            cancelled_name = cancelled_vendor['detail']['title'] if cancelled_vendor['detail'] else cancelled_vendor['request_name']
            cancelled_names.add(cancelled_name)
            for tag in tags:
                if tag:
                    combined_all_tags[tag].add(('cancelled', cancelled_name))
    
    # Sort key function for numeric table numbers
    def sort_key_for_index(entry):
        entry_type, entry_value = entry
        
        # Active vendors (numbers) come first, sorted numerically
        if entry_type == 'number':
            try:
                # Extract leading number from strings like "19-21" or "3 & 4"
                match = re.match(r'\s*(\d+)', str(entry_value))
                if match:
                    return (0, int(match.group(1)), str(entry_value))
            except (ValueError, AttributeError):
                pass
            return (0, float('inf'), str(entry_value))
        
        # Cancelled vendors come second, alphabetically
        return (1, 0, str(entry_value))
    
    # Render the index
    for tag in sorted(combined_all_tags.keys()):
        vendors = sorted(combined_all_tags[tag], key=sort_key_for_index)
        if not vendors:
            continue

        document.add_heading(tag, level=3)
        for entry_type, entry_value in vendors:
            if entry_type == 'number':
                # Active vendor: use number without bullet
                para = document.add_paragraph(f"{entry_value}. {num_to_name.get(entry_value, '')}")
            else:
                # Cancelled vendor: use bullet point
                para = document.add_paragraph(f"{entry_value}", style='List Bullet')
            
            # Set consistent spacing for all index entries
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(6)

    document.save(output_path)

def build_vendor_detail_index(vendor_details):
    """Build normalized lookup keys for vendor detail records."""
    detail_index = {}

    for detail_name, detail_info in vendor_details.items():
        title = detail_info.get('title') or detail_name
        filename = detail_info.get('filename') or ''

        keys = {
            normalize_name(detail_name),
            normalize_name(title),
        }

        if filename:
            keys.add(normalize_name(filename))
            keys.add(normalize_name(os.path.splitext(filename)[0]))
            stem = os.path.splitext(filename)[0]
            if '-' in stem:
                keys.add(normalize_name(stem.split('-', 3)[-1]))

        for key in keys:
            if key:
                detail_index[key] = detail_info

    return detail_index

def resolve_vendor_detail(request_name, vendor_details, vendor_filename_map=None, detail_index=None):
    """Resolve a vendor detail record using filename map first, then normalized lookups."""
    if detail_index is None:
        detail_index = build_vendor_detail_index(vendor_details)

    request_norm = normalize_name(request_name)

    if vendor_filename_map:
        mapped = vendor_filename_map.get(request_norm)
        if mapped:
            filename = mapped.get('filename', '')
            if filename:
                filename_norm = normalize_name(filename)
                filename_stem_norm = normalize_name(os.path.splitext(filename)[0])
                filename_tail_norm = normalize_name(os.path.splitext(filename)[0].split('-', 3)[-1])

                if filename_norm in detail_index:
                    return detail_index[filename_norm]
                if filename_stem_norm in detail_index:
                    return detail_index[filename_stem_norm]
                if filename_tail_norm in detail_index:
                    return detail_index[filename_tail_norm]

                for detail_info in vendor_details.values():
                    detail_filename = detail_info.get('filename', '')
                    if detail_filename and (
                        detail_filename == filename or
                        os.path.basename(detail_filename) == os.path.basename(filename)
                    ):
                        return detail_info

    if request_norm in detail_index:
        return detail_index[request_norm]

    for detail_name, detail_info in vendor_details.items():
        if names_match(request_name, detail_name) or names_match(request_name, detail_info.get('title', '')):
            return detail_info

    return None

def generate_vendor_directory(vendor_requests, vendor_details, output_path, output_missing_path=None, vendor_map_path=None, vendor_filename_map=None, cancelled_names=None):
    """Generate the vendor directory markdown file."""
    
    # Match vendor requests with details
    vendors_combined = []
    
    # Normalize cancelled names
    cancelled_norm = set()
    if cancelled_names:
        for cn in cancelled_names:
            cancelled_norm.add(normalize_name(cn))

    # Build list while applying manual filename/number overrides and skipping placeholders
    detail_index = build_vendor_detail_index(vendor_details)
    for req_name, req_num in vendor_requests.items():
        if not req_name or req_name.strip().lower() == 'empty':
            continue

        # apply manual overrides from vendor_filename_map if provided
        detail = resolve_vendor_detail(req_name, vendor_details, vendor_filename_map, detail_index)
        mapped = vendor_filename_map.get(normalize_name(req_name)) if vendor_filename_map else None

        # override table number if mapping provides
        if mapped and mapped.get('table_override'):
            use_num = mapped.get('table_override')
        else:
            # default PPL First correction: ensure PPL First uses 63
            if normalize_name(req_name) == normalize_name('PPL First'):
                use_num = '63'
            else:
                use_num = req_num

        vendors_combined.append({
            'number': use_num,
            'request_name': req_name,
            'detail': detail,
            'is_cancelled': normalize_name(req_name) in cancelled_norm
        })
    
    def sort_key_number(val):
        # Try to parse an integer from the start of the identifier (handles '19-21', '3 & 4')
        try:
            s = str(val)
            # grab leading digits
            import re
            m = re.match(r"\s*(\d+)", s)
            if m:
                return (0, int(m.group(1)))
        except:
            pass
        # fallback: sort as string after numeric ones
        return (1, str(val))

    # sort and ensure cancelled vendors are placed at the end
    vendors_combined.sort(key=lambda x: sort_key_number(x['number']))
    active = [v for v in vendors_combined if not v.get('is_cancelled')]
    cancelled = [v for v in vendors_combined if v.get('is_cancelled')]
    vendors_combined = active + cancelled
    cancelled_lookup = {normalize_name(v['request_name']): v for v in cancelled}

    def detail_for_cancelled_name(cancelled_name):
        cancelled_norm_name = normalize_name(cancelled_name)
        if cancelled_norm_name in cancelled_lookup:
            return cancelled_lookup[cancelled_norm_name]

        detail_info = resolve_vendor_detail(cancelled_name, vendor_details, vendor_filename_map, detail_index)
        if detail_info:
            return {
                'number': None,
                'request_name': cancelled_name,
                'detail': detail_info,
                'is_cancelled': True,
            }

        return {
            'number': None,
            'request_name': cancelled_name,
            'detail': None,
            'is_cancelled': True,
        }

    cancelled_entries = []
    if cancelled_names:
        for cancelled_name in cancelled_names:
            cancelled_entries.append(detail_for_cancelled_name(cancelled_name))
    else:
        cancelled_entries = cancelled
    
    # Collect all categories and tags and build number->name map
    all_categories = defaultdict(set)
    all_tags = defaultdict(set)
    num_to_name = {}
    # canonical age groups used across the site
    AGE_GROUPS = [
        'Ages Birth-3',
        'Early Intervention (Ages 3-5)',
        'Elementary (Grades K-6)',
        'Secondary (Grades 7-12)',
        'Post Secondary (High School and Beyond)'
    ]
    
    # Use module-level `normalize_name` and `names_match` helpers
    
    for v in vendors_combined:
        # Map table number to display name (use detail title if available)
        display_name = v['detail']['title'] if v['detail'] else v['request_name']
        num_to_name[v['number']] = display_name

        if v['detail']:
            cats = parse_categories(v['detail']['categories'])
            for cat in cats:
                if cat:
                    all_categories[cat].add(v['number'])
            
            tags = parse_tags(v['detail']['tags'])
            for tag in tags:
                if tag:
                    all_tags[tag].add(v['number'])
    
    # Generate markdown content
    repo_root = os.path.dirname(os.path.abspath(__file__))
    assets_dir = os.path.join(repo_root, 'assets')
    rel_assets = os.path.relpath(assets_dir, os.path.dirname(os.path.abspath(output_path)))
    rel_assets_href = rel_assets.replace(os.sep, '/')

    md_content = build_vendor_directory_banner_markdown(f"{rel_assets_href}/img/logo.png") + """# Disability Community Resource Fair 2026 - Vendor Directory

This is the comprehensive vendor directory for the 2026 Disability Community Resource Fair. Vendors are listed by table number, with an index organized by service categories and age groups served.

---

## Vendor Listings by Table Number

"""
    
    for v in active:
        title = v['detail']['title'] if v['detail'] else v['request_name']
        md_content += f"\n### {v['number']}. {title}\n\n"

        if v['detail']:
            # Add categories as inline tags
            cats = parse_categories(v['detail']['categories'])
            if cats:
                cats_str = ' '.join([f'[{c}]' for c in cats if c])
                md_content += f"Services: {cats_str}\n\n"

            # Add tags (age groups) as inline tags
            tags = parse_tags(v['detail']['tags'])
            if tags:
                # show 'All ages' if vendor covers all known age groups
                tag_set = set(tags)
                if all(a in tag_set for a in AGE_GROUPS):
                    md_content += f"Age Groups: All ages\n\n"
                else:
                    tags_str = ' '.join([f'[{t}]' for t in tags if t])
                    md_content += f"Age Groups: {tags_str}\n\n"

            content = v['detail']['content'].strip()
            if content:
                md_content += f"{content}\n\n"

    if cancelled_entries:
        md_content += "\n---\n\n## Unable to Attend\n\n"
        for v in cancelled_entries:
            title = v['detail']['title'] if v['detail'] else v['request_name']
            md_content += f"\n### {title}\n\n"
            if v['detail']:
                cats = parse_categories(v['detail']['categories'])
                if cats:
                    cats_str = ' '.join([f'[{c}]' for c in cats if c])
                    md_content += f"Services: {cats_str}\n\n"

                tags = parse_tags(v['detail']['tags'])
                if tags:
                    tag_set = set(tags)
                    if all(a in tag_set for a in AGE_GROUPS):
                        md_content += f"Age Groups: All ages\n\n"
                    else:
                        tags_str = ' '.join([f'[{t}]' for t in tags if t])
                        md_content += f"Age Groups: {tags_str}\n\n"

                content = v['detail']['content'].strip()
                if content:
                    md_content += f"{content}\n\n"
            else:
                md_content += "Description unavailable.\n\n"
    
    # Add index section (only Age Group index, two-column lists)
    md_content += "\n---\n\n## Index by Age Group Served\n\n"

    for tag in sorted(all_tags.keys()):
        vendors = sorted(all_tags[tag])
        if vendors:
            md_content += f"\n### {tag}\n"
            entries = [f"{n}. {num_to_name.get(n, '')}" for n in vendors]
            # two-column markdown table
            md_content += "\n| | |\n|---|---|\n"
            left = entries[0::2]
            right = entries[1::2]
            for i in range(max(len(left), len(right))):
                L = left[i] if i < len(left) else ''
                R = right[i] if i < len(right) else ''
                md_content += f"| {L} | {R} |\n"
    
    # Write the directory to file
    if str(output_path).lower().endswith('.docx'):
        write_vendor_directory_docx(
            output_path,
            active,
            cancelled_entries,
            all_tags,
            num_to_name,
            AGE_GROUPS,
        )
    elif str(output_path).lower().endswith('.html'):
        # Compute relative path from output file to repo assets so links work when opened via file://
        # Use the site's CSS and fonts so the vendor directory matches the main site
        html_parts = ["<!doctype html>", "<html lang=\"en\">", "<head>", "<meta charset=\"utf-8\">",
                      "<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">",
                      f"<title>Disability Community Resource Fair 2026 - Vendor Directory</title>",
                      f'<link rel="stylesheet" href="{rel_assets_href}/css/bootstrap.min.css">',
                      f'<link rel="stylesheet" href="{rel_assets_href}/css/academicons.min.css">',
                      f'<link rel="stylesheet" href="{rel_assets_href}/css/scholar-icons.css">',
                  f'<link rel="stylesheet" href="{rel_assets_href}/css/main.css">',
                  "<style>",
                  ".badge{display:inline-block;padding:2px 6px;margin:2px;border-radius:4px;font-size:0.82em}",
                  ".badge-service{background:#e6f4ff;color:#023d66}",
                  ".badge-age{background:#fff4e6;color:#4a6e1b}",
                  ".badge-allages{background:#f0e6ff;color:#4a1766}",
                      ".badge-status{background:#eef2f7;color:#4b5563}",
                  "</style>",
                  "</head>", "<body>"]

        html_parts.extend(build_vendor_directory_banner_html(f"{rel_assets_href}/img/logo.png"))
        html_parts.append("<h1>Disability Community Resource Fair 2026 - Vendor Directory</h1>")
        html_parts.append("<p>Vendors are listed by table number, with services and full descriptions.</p>")
        html_parts.append("<hr>")
        html_parts.append("<section id=\"listings\">")

        for v in active:
            title = v['detail']['title'] if v['detail'] else v['request_name']
            html_parts.append(f"<h3>{v['number']}. {title}</h3>")
            if v['detail']:
                cats = parse_categories(v['detail']['categories'])
                tags = parse_tags(v['detail']['tags'])
                badge_parts = []
                if cats:
                    badge_parts.extend([f"<span class=\"badge badge-service\">{c}</span>" for c in cats if c])
                if tags:
                    tag_set = set(tags)
                    if all(a in tag_set for a in AGE_GROUPS):
                        badge_parts.append("<span class=\"badge badge-age badge-allages\">All ages</span>")
                    else:
                        badge_parts.extend([f"<span class=\"badge badge-age\">{t}</span>" for t in tags if t])
                if badge_parts:
                    # keep service and age tags on one line
                    html_parts.append(f"<p>{' '.join(badge_parts)}</p>")
                content = v['detail']['content'].strip()
                if content:
                    # simple paragraph wrapping; preserve newlines
                    paragraphs = [f"<p>{p.strip()}</p>" for p in content.split('\n\n') if p.strip()]
                    html_parts.extend(paragraphs)

        if cancelled_entries:
            html_parts.append("<hr>")
            html_parts.append("<section id=\"unable-to-attend\">")
            html_parts.append("<h2>Unable to Attend</h2>")
            for v in cancelled_entries:
                title = v['detail']['title'] if v['detail'] else v['request_name']
                html_parts.append(f"<h3>{title}</h3>")
                if v['detail']:
                    cats = parse_categories(v['detail']['categories'])
                    tags = parse_tags(v['detail']['tags'])
                    badge_parts = []
                    if cats:
                        badge_parts.extend([f"<span class=\"badge badge-service\">{c}</span>" for c in cats if c])
                    if tags:
                        tag_set = set(tags)
                        if all(a in tag_set for a in AGE_GROUPS):
                            badge_parts.append("<span class=\"badge badge-age badge-allages\">All ages</span>")
                        else:
                            badge_parts.extend([f"<span class=\"badge badge-age\">{t}</span>" for t in tags if t])
                    if badge_parts:
                        html_parts.append(f"<p>{' '.join(badge_parts)}</p>")
                    content = v['detail']['content'].strip()
                    if content:
                        paragraphs = [f"<p>{p.strip()}</p>" for p in content.split('\n\n') if p.strip()]
                        html_parts.extend(paragraphs)
                else:
                    html_parts.append("<p>Description unavailable.</p>")
            html_parts.append("</section>")

        html_parts.append("</section>")

        # Index by age group only. Render as two-column responsive grid
        html_parts.append("<hr>")
        html_parts.append("<section id=\"index-age\">")
        html_parts.append("<h2>Index by Age Group Served</h2>")
        html_parts.append('<style>.age-grid{display:grid;grid-template-columns:1fr 1fr;gap:8px}.age-group{margin-bottom:12px}.age-group h3{margin:4px 0}</style>')
        html_parts.append('<div class="age-grid">')
        for tag in sorted(all_tags.keys()):
            vendors = sorted(all_tags[tag])
            if vendors:
                html_parts.append(f"<div class=\"age-group\"><h3>{tag}</h3>")
                html_parts.append('<ul>')
                for n in vendors:
                    html_parts.append(f"<li>{n}. {num_to_name.get(n, '')}</li>")
                html_parts.append('</ul></div>')
        html_parts.append('</div>')
        html_parts.append("</section>")

        html_parts.append("</body></html>")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_parts))
    else:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)

    # Generate missing vendors report (two-way):
    # 1) Vendors present in requests but without a markdown file
    # 2) Vendors present in _posts but not referenced in the requests/logistics
    missing = [v for v in vendors_combined if not v['detail']]

    vendor_map_conflicts = []
    if vendor_map_path:
        vendor_map_entries = read_vendor_map_entries(vendor_map_path)
        logistics_entries = []

        for req_name, req_num in vendor_requests.items():
            logistics_entries.append({
                'display_name': req_name,
                'normalized_name': normalize_name(req_name),
                'number_values': parse_table_number_set(req_num),
            })

        for map_entry in vendor_map_entries:
            matching_logistics = [entry for entry in logistics_entries if names_match(map_entry['display_name'], entry['display_name'])]
            best_match = None
            best_score = 0.0

            if matching_logistics:
                import difflib

                for entry in matching_logistics:
                    score = difflib.SequenceMatcher(a=normalize_name(map_entry['display_name']), b=normalize_name(entry['display_name'])).ratio()
                    if score > best_score:
                        best_score = score
                        best_match = entry

            if best_match is not None:
                if map_entry['number_values'] != best_match['number_values']:
                    vendor_map_conflicts.append({
                        'map_name': map_entry['display_name'],
                        'map_numbers': format_table_number_set(map_entry['number_values']),
                        'logistics_name': best_match['display_name'],
                        'logistics_numbers': format_table_number_set(best_match['number_values']),
                        'conflict_type': 'vendor matched, table numbers differ',
                    })
                continue

            overlapping_logistics = [entry for entry in logistics_entries if map_entry['number_values'] & entry['number_values']]
            if overlapping_logistics:
                overlap_entry = sorted(overlapping_logistics, key=lambda entry: entry['display_name'])[0]
                vendor_map_conflicts.append({
                    'map_name': map_entry['display_name'],
                    'map_numbers': format_table_number_set(map_entry['number_values']),
                    'logistics_name': overlap_entry['display_name'],
                    'logistics_numbers': format_table_number_set(overlap_entry['number_values']),
                    'conflict_type': 'table number overlaps a different vendor',
                })

    # Build normalized name maps for comparison
    logistics_names = set()
    logistics_norm_to_original = {}
    for req_name in vendor_requests.keys():
        n = normalize_name(req_name)
        logistics_names.add(n)
        logistics_norm_to_original[n] = req_name

    posts_names = set()
    posts_norm_to_original = {}
    for pd_name, pd in vendor_details.items():
        # prefer the title if present
        display = pd.get('title') or pd_name
        n = normalize_name(display)
        posts_names.add(n)
        posts_norm_to_original[n] = (display, pd.get('filename'))

    matched_posts_norm = set()
    for p_norm in posts_names:
        display, _filename = posts_norm_to_original.get(p_norm, ('', ''))
        if any(names_match(req_name, display) for req_name in vendor_requests.keys()):
            matched_posts_norm.add(p_norm)

    posts_not_in_logistics_norm = sorted([n for n in posts_names if n not in matched_posts_norm])

    # Attempt fuzzy reconciliation for missing logistics entries vs posts
    import difflib
    reconciled = []
    unresolved_missing = []
    for m in missing:
        req = m['request_name']
        req_norm = normalize_name(req)
        # find best match in posts
        best = None
        best_score = 0.0
        for p_norm in posts_names:
            score = difflib.SequenceMatcher(a=req_norm, b=p_norm).ratio()
            if score > best_score:
                best_score = score
                best = p_norm

        # consider it reconciled if similarity is high or one name is substring of the other
        if best and (best_score >= 0.72 or best in req_norm or req_norm in best):
            display, filename = posts_norm_to_original.get(best, ('', ''))
            reconciled.append({'request_number': m['number'], 'request_name': req, 'matched_post': display, 'filename': filename, 'score': round(best_score, 2)})
        else:
            unresolved_missing.append(m)

    # Replace missing with unresolved_missing for the report
    missing = unresolved_missing

    if output_missing_path is not None:
        # If the user requested an HTML report, produce a printer-friendly HTML
        if str(output_missing_path).lower().endswith('.html'):
            html = []
            html.append('<!doctype html>')
            html.append('<html lang="en">')
            html.append('<head>')
            html.append('<meta charset="utf-8">')
            html.append('<meta name="viewport" content="width=device-width,initial-scale=1">')
            html.append('<title>Missing Vendor Detail Files</title>')
            html.append('<style>')
            html.append('body{font-family:Arial,Helvetica,sans-serif;margin:20px;color:#111}');
            html.append('h1,h2{margin:0 0 10px 0}');
            html.append('table{border-collapse:collapse;width:100%;margin-top:8px}');
            html.append('th,td{border:1px solid #333;padding:6px;text-align:left}');
            html.append('@media print{body{margin:6mm} a[href]:after{content:""}}');
            html.append('</style>')
            html.append('</head>')
            html.append('<body>')
            html.append('<h1>Missing Vendor Detail Files</h1>')
            html.append('<p>This report lists two sets of mismatches between the fair logistics sheet and the <code>_posts</code> vendor detail files.</p>')

            # Section A
            html.append('<h2>A. In logistics but missing <code>_posts</code> file</h2>')
            html.append('<p>Vendors present in fair logistics but missing a corresponding <code>_posts</code> markdown file.</p>')
            html.append('<table>')
            html.append('<thead><tr><th style="width:15%">Table Number</th><th>Vendor Name</th></tr></thead>')
            html.append('<tbody>')
            for m in missing:
                if not m['request_name'] or m['request_name'].strip().lower() == 'empty':
                    continue
                html.append(f"<tr><td>{m['number']}</td><td>{m['request_name']}</td></tr>")
            html.append('</tbody></table>')

            # Section B
            html.append('<h2>B. In <code>_posts</code> but not in logistics</h2>')
            html.append('<p>Vendors that have a <code>_posts</code> detail file but were not found in the fair logistics vendor requests.</p>')
            html.append('<table>')
            html.append('<thead><tr><th>Vendor Filename</th><th>Vendor Name</th></tr></thead>')
            html.append('<tbody>')
            for n in posts_not_in_logistics_norm:
                display, filename = posts_norm_to_original.get(n, ('', ''))
                if not display or display.strip().lower() == 'empty':
                    continue
                html.append(f"<tr><td>{filename or ''}</td><td>{display}</td></tr>")
            html.append('</tbody></table>')

            if vendor_map_conflicts:
                html.append('<h2>C. In vendor_map but number conflicts with logistics</h2>')
                html.append('<p>Vendors that appear in both workbooks but have different table numbers, or share a table number with a different vendor.</p>')
                html.append('<table>')
                html.append('<thead><tr><th>Vendor Map Name</th><th>Vendor Map Number(s)</th><th>Fair Logistics Name</th><th>Fair Logistics Number(s)</th><th>Conflict Type</th></tr></thead>')
                html.append('<tbody>')
                for conflict in vendor_map_conflicts:
                    html.append(
                        '<tr>'
                        f"<td>{conflict['map_name']}</td>"
                        f"<td>{conflict['map_numbers']}</td>"
                        f"<td>{conflict['logistics_name']}</td>"
                        f"<td>{conflict['logistics_numbers']}</td>"
                        f"<td>{conflict['conflict_type']}</td>"
                        '</tr>'
                    )
                html.append('</tbody></table>')

            # Print-on-load script
            html.append('<script>window.addEventListener("load", function(){window.print();});</script>')
            html.append('</body></html>')

            with open(output_missing_path, 'w', encoding='utf-8') as r:
                r.write('\n'.join(html))
        else:
            # Fallback: write markdown as before
            rpt_lines = ["# Missing Vendor Detail Files\n\n",
                         "This report lists two sets of mismatches between the fair logistics sheet and the `_posts` vendor detail files.\n\n",
                         "## A. In logistics but missing `_posts` file\n\n",
                         "Vendors present in fair logistics but missing a corresponding `_posts` markdown file.\n\n",
                         "| Table Number | Vendor Name |\n",
                         "|---:|---|\n"]
            for m in missing:
                if not m['request_name'] or m['request_name'].strip().lower() == 'empty':
                    continue
                rpt_lines.append(f"| {m['number']} | {m['request_name']} |\n")

            rpt_lines.extend(["\n## B. In `_posts` but not in logistics\n\n",
                              "Vendors that have a `_posts` detail file but were not found in the fair logistics vendor requests.\n\n",
                              "| Vendor Filename | Vendor Name |\n",
                              "|---|---|\n"])

            for n in posts_not_in_logistics_norm:
                display, filename = posts_norm_to_original.get(n, ('', ''))
                if not display or display.strip().lower() == 'empty':
                    continue
                rpt_lines.append(f"| {filename or ''} | {display} |\n")

            if vendor_map_conflicts:
                rpt_lines.extend(["\n## C. In vendor_map but number conflicts with logistics\n\n",
                                  "Vendors that appear in both workbooks but have different table numbers, or share a table number with a different vendor.\n\n",
                                  "| Vendor Map Name | Vendor Map Number(s) | Fair Logistics Name | Fair Logistics Number(s) | Conflict Type |\n",
                                  "|---|---|---|---|---|\n"])

                for conflict in vendor_map_conflicts:
                    rpt_lines.append(
                        f"| {conflict['map_name']} | {conflict['map_numbers']} | {conflict['logistics_name']} | {conflict['logistics_numbers']} | {conflict['conflict_type']} |\n"
                    )

            with open(output_missing_path, 'w', encoding='utf-8') as r:
                r.writelines(rpt_lines)

    return {
        'total_vendors': len(vendors_combined),
        'service_categories': len(all_categories),
        'age_groups': len(all_tags),
        'file_size': len(md_content),
        'missing_count': len(missing),
        'missing_list': missing,
        'posts_not_in_logistics_count': len(posts_not_in_logistics_norm),
        'posts_not_in_logistics_list': posts_not_in_logistics_norm,
        'vendor_map_conflicts_count': len(vendor_map_conflicts),
        'vendor_map_conflicts_list': vendor_map_conflicts,
    }

def main():
    """Main execution function."""
    import sys

    parser = argparse.ArgumentParser(description='Generate the 2026 Disability Community Resource Fair vendor directory.')
    parser.add_argument('--format', choices=['docx', 'html', 'md'], default='docx', help='Output format for the vendor directory.')
    args = parser.parse_args()
    
    # Get base directory (parent of this script)
    base_dir = os.path.dirname(os.path.abspath(__file__))
    data_dir = os.path.join(base_dir, '_data')
    posts_dir = os.path.join(base_dir, '_posts')
    
    # File paths
    fair_logistics_path = os.path.join(data_dir, 'fair_logistics.xlsx')
    vendor_requests_csv = os.path.join(data_dir, '2026_vendor_requests.csv')
    vendor_directory_md = os.path.join(data_dir, f'2026_vendor_directory.{args.format}')
    vendor_missing_report = os.path.join(data_dir, '2026_missing_vendors.md')
    vendor_map_path = os.path.join(data_dir, 'vendor_map.xlsx')
    vendor_filename_map_path = os.path.join(data_dir, 'vendor_filename_map.csv')
    
    print("=" * 70)
    print("2026 Disability Community Resource Fair - Vendor Directory Generator")
    print("=" * 70)
    
    # Step 1: Extract vendors from Excel
    print("\n[1/4] Extracting vendors from fair_logistics.xlsx...")
    try:
        count = extract_vendors_from_excel(fair_logistics_path, vendor_requests_csv)
        print(f"      ✓ Extracted {count} rows to {vendor_requests_csv}")
    except Exception as e:
        print(f"      ✗ Error: {e}")
        sys.exit(1)
    
    # Step 2: Read vendor requests
    print("\n[2/4] Reading vendor requests from CSV...")
    try:
        vendor_requests = read_vendor_requests(vendor_requests_csv)
        print(f"      ✓ Loaded {len(vendor_requests)} vendor requests")
    except Exception as e:
        print(f"      ✗ Error: {e}")
        sys.exit(1)
    
    # Step 3: Read vendor details
    print("\n[3/4] Reading vendor details from markdown files...")
    try:
        vendor_details = read_vendor_details(posts_dir)
        print(f"      ✓ Loaded {len(vendor_details)} vendor detail files")
    except Exception as e:
        print(f"      ✗ Error: {e}")
        sys.exit(1)
    
    # Step 4: Generate vendor directory
    print("\n[4/4] Generating vendor directory...")
    try:
        vendor_filename_map = read_vendor_filename_map(vendor_filename_map_path)
        cancelled = [
            "Luke5Adventures",
            "Carrita Counseling",
            "Keck Health",
            "DVB Financial",
            "Technology Assisted Children’s Home Program (TACHP)",
        ]

        stats = generate_vendor_directory(
            vendor_requests,
            vendor_details,
            vendor_directory_md,
            output_missing_path=vendor_missing_report,
            vendor_map_path=vendor_map_path,
            vendor_filename_map=vendor_filename_map,
            cancelled_names=cancelled,
        )
        print(f"      ✓ Generated {vendor_directory_md}")
        print(f"\n      Statistics:")
        print(f"        - Vendors: {stats['total_vendors']}")
        print(f"        - Service Categories: {stats['service_categories']}")
        print(f"        - Age Groups: {stats['age_groups']}")
        print(f"        - File Size: {stats['file_size']:,} bytes")
        print(f"        - Missing vendor detail files: {stats.get('missing_count', 0)}")
        print(f"        - Vendor map conflicts: {stats.get('vendor_map_conflicts_count', 0)}")
    except Exception as e:
        print(f"      ✗ Error: {e}")
        sys.exit(1)
    
    print("\n" + "=" * 70)
    print("✓ Vendor directory generation complete!")
    print("=" * 70)

if __name__ == '__main__':
    main()
